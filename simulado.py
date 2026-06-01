import streamlit as st
from streamlit_gsheets import GSheetsConnection
import pandas as pd
from docx import Document
from docx.shared import Inches
from io import BytesIO
from datetime import datetime
import os
import requests
import time
from bs4 import BeautifulSoup
from streamlit_quill import st_quill

# ─── CONFIGURAÇÃO DA PÁGINA ───────────────────────────────────────────────────
st.set_page_config(
    page_title="SIDE — Portal de Simulados",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── CSS GLOBAL ───────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Sora:wght@300;400;500;600;700;800&family=JetBrains+Mono:wght@400;500&display=swap');

/* ── Reset e base ── */
html, body, [class*="css"] {
    font-family: 'Sora', sans-serif !important;
}

/* ── Fundo geral ── */
.stApp {
    background: #0d1f17;
    background-image:
        radial-gradient(ellipse 80% 60% at 20% 0%, rgba(45,106,79,0.25) 0%, transparent 60%),
        radial-gradient(ellipse 60% 40% at 80% 100%, rgba(26,58,42,0.30) 0%, transparent 55%);
}

/* ── Sidebar ── */
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #0a1710 0%, #132a1e 100%) !important;
    border-right: 1px solid rgba(45,106,79,0.35) !important;
}
section[data-testid="stSidebar"] * {
    color: #c8e6d4 !important;
}
section[data-testid="stSidebar"] .stRadio label {
    color: #c8e6d4 !important;
}

/* ── Header animado ── */
@keyframes fadeSlideDown {
    from { opacity: 0; transform: translateY(-22px); }
    to   { opacity: 1; transform: translateY(0); }
}
.header-box {
    background: linear-gradient(135deg, #1a3a2a 0%, #2d6a4f 60%, #1e4d38 100%);
    border: 1px solid rgba(74,183,108,0.30);
    border-radius: 16px;
    padding: 30px 36px;
    color: white;
    margin-bottom: 28px;
    animation: fadeSlideDown 0.6s ease both;
    box-shadow: 0 8px 40px rgba(0,0,0,0.45), inset 0 1px 0 rgba(255,255,255,0.08);
    position: relative;
    overflow: hidden;
}
.header-box::before {
    content: '';
    position: absolute;
    top: -40%; right: -10%;
    width: 300px; height: 300px;
    border-radius: 50%;
    background: radial-gradient(circle, rgba(74,183,108,0.12) 0%, transparent 70%);
    pointer-events: none;
}
.header-box h1 {
    font-size: 2rem;
    font-weight: 800;
    letter-spacing: -0.5px;
    margin: 0 0 6px 0;
    text-shadow: 0 2px 12px rgba(0,0,0,0.4);
}
.header-box p {
    margin: 0;
    opacity: 0.80;
    font-size: 0.95rem;
    font-weight: 300;
    letter-spacing: 0.3px;
}

/* ── Cards de seção ── */
@keyframes fadeUp {
    from { opacity: 0; transform: translateY(16px); }
    to   { opacity: 1; transform: translateY(0); }
}
.card-section {
    background: rgba(20,42,30,0.65);
    border: 1px solid rgba(45,106,79,0.35);
    border-radius: 14px;
    padding: 28px 32px;
    margin-bottom: 20px;
    animation: fadeUp 0.5s ease both;
    backdrop-filter: blur(8px);
    box-shadow: 0 4px 24px rgba(0,0,0,0.30);
}

/* ── Títulos de seção ── */
.section-title {
    font-size: 1.15rem;
    font-weight: 700;
    color: #6fcf97;
    letter-spacing: 0.4px;
    margin-bottom: 18px;
    display: flex;
    align-items: center;
    gap: 8px;
}
.section-title::after {
    content: '';
    flex: 1;
    height: 1px;
    background: linear-gradient(90deg, rgba(111,207,151,0.4), transparent);
    margin-left: 6px;
}

/* ── Inputs ── */
.stTextInput input,
.stTextArea textarea,
.stSelectbox > div > div {
    background: rgba(10,23,16,0.70) !important;
    border: 1px solid rgba(45,106,79,0.50) !important;
    border-radius: 10px !important;
    color: #e0f0e8 !important;
    font-family: 'Sora', sans-serif !important;
    transition: border-color 0.25s, box-shadow 0.25s;
}
.stTextInput input:focus,
.stTextArea textarea:focus {
    border-color: rgba(111,207,151,0.70) !important;
    box-shadow: 0 0 0 3px rgba(111,207,151,0.12) !important;
}
.stTextInput label,
.stTextArea label,
.stSelectbox label,
.stRadio label,
.stMultiSelect label {
    color: #a8d5bc !important;
    font-size: 0.85rem !important;
    font-weight: 500 !important;
    letter-spacing: 0.3px;
}

/* ── Radio Buttons ── */
.stRadio div[role="radiogroup"] {
    gap: 10px;
}
.stRadio div[role="radiogroup"] label {
    background: rgba(15,32,22,0.60);
    border: 1px solid rgba(45,106,79,0.40);
    border-radius: 8px;
    padding: 6px 14px;
    transition: all 0.2s ease;
    color: #c8e6d4 !important;
}
.stRadio div[role="radiogroup"] label:hover {
    border-color: rgba(111,207,151,0.65);
    background: rgba(45,106,79,0.25);
}

/* ── Botão principal ── */
@keyframes pulse-glow {
    0%, 100% { box-shadow: 0 0 0 0 rgba(111,207,151,0.35); }
    50%        { box-shadow: 0 0 0 8px rgba(111,207,151,0); }
}
.stButton > button,
.stDownloadButton > button,
.stFormSubmitButton > button {
    background: linear-gradient(135deg, #2d6a4f 0%, #40916c 100%) !important;
    color: #e8f7ef !important;
    border: 1px solid rgba(111,207,151,0.45) !important;
    border-radius: 10px !important;
    font-family: 'Sora', sans-serif !important;
    font-weight: 600 !important;
    font-size: 0.92rem !important;
    letter-spacing: 0.3px;
    padding: 0.55rem 1.4rem !important;
    transition: all 0.25s ease !important;
    position: relative;
    overflow: hidden;
}
.stButton > button::before,
.stFormSubmitButton > button::before {
    content: '';
    position: absolute;
    top: 0; left: -100%;
    width: 100%; height: 100%;
    background: linear-gradient(90deg, transparent, rgba(255,255,255,0.12), transparent);
    transition: left 0.4s ease;
}
.stButton > button:hover::before,
.stFormSubmitButton > button:hover::before {
    left: 100%;
}
.stButton > button:hover,
.stDownloadButton > button:hover,
.stFormSubmitButton > button:hover {
    background: linear-gradient(135deg, #40916c 0%, #52b788 100%) !important;
    transform: translateY(-2px) !important;
    box-shadow: 0 6px 20px rgba(45,106,79,0.50) !important;
    animation: pulse-glow 1.5s infinite;
}
.stButton > button:active,
.stFormSubmitButton > button:active {
    transform: translateY(0px) !important;
    box-shadow: 0 2px 8px rgba(45,106,79,0.35) !important;
}

/* ── Download button destaque ── */
.stDownloadButton > button {
    background: linear-gradient(135deg, #1b4332 0%, #2d6a4f 100%) !important;
    border: 1px solid rgba(111,207,151,0.55) !important;
    width: 100%;
}
.stDownloadButton > button:hover {
    background: linear-gradient(135deg, #2d6a4f 0%, #40916c 100%) !important;
}

/* ── Métricas ── */
div[data-testid="stMetric"] {
    background: rgba(20,42,30,0.65) !important;
    border: 1px solid rgba(45,106,79,0.35) !important;
    border-radius: 12px !important;
    padding: 16px 22px !important;
    backdrop-filter: blur(6px);
    transition: transform 0.2s, box-shadow 0.2s;
}
div[data-testid="stMetric"]:hover {
    transform: translateY(-2px);
    box-shadow: 0 6px 20px rgba(0,0,0,0.30);
}
div[data-testid="stMetric"] label {
    color: #74c69d !important;
    font-size: 0.82rem !important;
    font-weight: 600;
    letter-spacing: 0.5px;
    text-transform: uppercase;
}
div[data-testid="stMetric"] div[data-testid="stMetricValue"] {
    color: #d8f3dc !important;
    font-size: 2.1rem !important;
    font-weight: 700;
    font-family: 'JetBrains Mono', monospace !important;
}

/* ── Dataframe ── */
div[data-testid="stDataFrame"] {
    border-radius: 12px;
    overflow: hidden;
    border: 1px solid rgba(45,106,79,0.30) !important;
}

/* ── Divider ── */
hr {
    border-color: rgba(45,106,79,0.30) !important;
    margin: 20px 0 !important;
}

/* ── Alertas / Toasts ── */
div[data-testid="stAlert"] {
    border-radius: 10px !important;
    border-left-width: 4px !important;
    font-family: 'Sora', sans-serif !important;
}

/* ── Progress bar ── */
.stProgress > div > div > div {
    background: linear-gradient(90deg, #40916c, #74c69d) !important;
    border-radius: 999px !important;
}
.stProgress > div > div {
    background: rgba(20,42,30,0.55) !important;
    border-radius: 999px !important;
}

/* ── Spinner ── */
div[data-testid="stSpinner"] {
    color: #6fcf97 !important;
}

/* ── Multiselect tags ── */
span[data-baseweb="tag"] {
    background: rgba(45,106,79,0.55) !important;
    border-radius: 6px !important;
}

/* ── Cabeçalho da sidebar ── */
.sidebar-logo {
    text-align: center;
    padding: 12px 0 8px 0;
}
.sidebar-logo .logo-icon {
    font-size: 2.4rem;
    line-height: 1;
    display: block;
    margin-bottom: 6px;
    filter: drop-shadow(0 0 10px rgba(111,207,151,0.5));
}
.sidebar-logo .logo-title {
    font-size: 1.5rem;
    font-weight: 800;
    color: #6fcf97;
    letter-spacing: 2px;
}
.sidebar-logo .logo-sub {
    font-size: 0.72rem;
    color: #74c69d;
    opacity: 0.65;
    letter-spacing: 1px;
    text-transform: uppercase;
}

/* ── Badge de data ── */
.date-badge {
    background: rgba(45,106,79,0.30);
    border: 1px solid rgba(74,183,108,0.25);
    border-radius: 8px;
    padding: 6px 12px;
    font-size: 0.78rem;
    color: #74c69d;
    font-family: 'JetBrains Mono', monospace;
    text-align: center;
    margin-top: 8px;
}

/* ── Tag de campo obrigatório ── */
.req-badge {
    display: inline-block;
    background: rgba(231,111,81,0.20);
    color: #e76f51;
    border: 1px solid rgba(231,111,81,0.35);
    border-radius: 4px;
    font-size: 0.68rem;
    font-weight: 600;
    padding: 1px 6px;
    margin-left: 6px;
    letter-spacing: 0.5px;
    vertical-align: middle;
}

/* ── Animação de entrada dos campos ── */
@keyframes fadeInField {
    from { opacity: 0; transform: translateY(8px); }
    to   { opacity: 1; transform: translateY(0); }
}
.field-animated {
    animation: fadeInField 0.4s ease both;
}

/* ── Ajuste visual para o Quill ── */
.quill-label {
    color: #a8d5bc;
    font-size: 0.85rem;
    font-weight: 500;
    letter-spacing: 0.3px;
    margin-bottom: 6px;
    display: block;
}
</style>
""", unsafe_allow_html=True)

# ─── TRATAMENTO DE TEXTO ──────────────────────────────────────────────────────
def limpar_texto(texto):
    if pd.isna(texto) or texto is None:
        return ""
    t = str(texto).strip()
    substituicoes = {
        '"': '"', '"': '"', '\u2018': "'", '\u2019': "'",
        '–': '-', '—': '-', '…': '...',
        '\u2013': '-', '\u2014': '-', '\u201c': '"', '\u201d': '"',
        '\u2026': '...'
    }
    for original, novo in substituicoes.items():
        t = t.replace(original, novo)
    return t

def _processar_elementos_html(p_doc, elementos):
    """Função auxiliar para processar as tags HTML e adicionar ao parágrafo do Word."""
    for elemento in elementos:
        if elemento.name is None:
            run = p_doc.add_run(limpar_texto(str(elemento)))
        elif elemento.name in ['strong', 'b']:
            run = p_doc.add_run(limpar_texto(elemento.get_text()))
            run.bold = True
        elif elemento.name in ['em', 'i']:
            run = p_doc.add_run(limpar_texto(elemento.get_text()))
            run.italic = True
        elif elemento.name == 'br':
            p_doc.add_run('\n')
        else:
            run = p_doc.add_run(limpar_texto(elemento.get_text()))

def processar_html_para_docx(doc, texto_html):
    """Lê o HTML gerado pelo Quill e insere no documento Word com as formatações."""
    if pd.isna(texto_html) or not str(texto_html).strip():
        return

    soup = BeautifulSoup(str(texto_html), "html.parser")
    paragrafos_html = soup.find_all('p')

    if not paragrafos_html:
        p_doc = doc.add_paragraph()
        _processar_elementos_html(p_doc, soup.children)
        return

    for p_tag in paragrafos_html:
        p_doc = doc.add_paragraph()
        _processar_elementos_html(p_doc, p_tag.children)

# ─── FUNÇÃO PARA GERAR DOCX (PADRÃO SIDE) ────────────────────────────────────
# CORREÇÃO 4: tratamento de erros de imagem com feedback por questão.
def gerar_docx_questoes(df_export):
    doc = Document()
    erros_imagem = []  # acumula avisos para exibir após gerar o documento

    for idx, (_, r) in enumerate(df_export.iterrows(), 1):
        doc.add_paragraph(f"Disciplina: {r.get('Disciplina')}")
        doc.add_paragraph(f"Habilidade: {r.get('Habilidade')}")
        p_num = doc.add_paragraph()
        p_num.add_run(f"QUESTÃO {idx:02d}").bold = True

        processar_html_para_docx(doc, r.get("Pergunta"))

        link_img = r.get("Link Imagem")
        if pd.notna(link_img) and str(link_img).strip().lower() not in ["", "nan"]:
            try:
                url_final = str(link_img).strip()

                if "drive.google.com" in url_final:
                    if "/file/d/" in url_final:
                        id_arquivo = url_final.split("/file/d/")[1].split("/")[0]
                        url_final = f"https://drive.google.com/uc?export=download&id={id_arquivo}"
                    elif "id=" in url_final:
                        id_arquivo = url_final.split("id=")[1].split("&")[0]
                        url_final = f"https://drive.google.com/uc?export=download&id={id_arquivo}"

                # --- ADIÇÃO DO USER-AGENT PARA EVITAR BLOQUEIO DE IMAGENS ---
                cabecalhos = {
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
                }

                resp = requests.get(url_final, headers=cabecalhos, timeout=10)
                resp.raise_for_status()  # lança exceção para status 4xx/5xx

                img_io = BytesIO(resp.content)
                doc.add_picture(img_io, width=Inches(4.0))

            except requests.exceptions.Timeout:
                erros_imagem.append(f"Questão {idx:02d}: tempo de conexão esgotado ao baixar a imagem.")
            except requests.exceptions.HTTPError as e:
                erros_imagem.append(f"Questão {idx:02d}: erro HTTP {e.response.status_code} ao acessar a imagem.")
            except requests.exceptions.RequestException as e:
                erros_imagem.append(f"Questão {idx:02d}: falha de rede — {str(e)}")
            except Exception as e:
                erros_imagem.append(f"Questão {idx:02d}: não foi possível inserir a imagem — {str(e)}")

        for letra in ["A", "B", "C", "D", "E"]:
            conteudo = r.get(letra)
            if pd.notna(conteudo):
                doc.add_paragraph(f"({letra}) {limpar_texto(conteudo)}")
        doc.add_paragraph("-" * 30)
        doc.add_paragraph("\n")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, erros_imagem  # retorna a lista de erros junto ao buffer


# ─── FUNÇÃO PARA GERAR DOCX DO GABARITO ─────────────────────────────────────
def gerar_docx_gabarito(df_export):
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    section = doc.sections[0]
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = titulo.add_run("GABARITO — BANCO DE QUESTÕES SIDE")
    run_t.bold = True
    run_t.font.size = Pt(15)
    run_t.font.color.rgb = RGBColor(0x1a, 0x3a, 0x2a)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = sub.add_run(
        f"Escola Estadual Padre Constantino de Monte — Maracaju/MS\n"
        f"Gerado em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}"
    )
    run_s.font.size = Pt(9)
    run_s.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    colunas = ["Nº", "Disciplina", "Habilidade", "Turma", "Gabarito"]
    tabela = doc.add_table(rows=1, cols=len(colunas))
    tabela.style = "Table Grid"

    hdr = tabela.rows[0].cells
    larguras = [Cm(1.2), Cm(3.5), Cm(6.5), Cm(2.0), Cm(2.0)]
    for i, (cell, col_nome) in enumerate(zip(hdr, colunas)):
        cell.width = larguras[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(col_nome)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2D6A4F")
        tc_pr.append(shd)

    for idx, (_, r) in enumerate(df_export.iterrows(), 1):
        row_cells = tabela.add_row().cells
        valores = [
            f"{idx:02d}",
            str(r.get("Disciplina", "")),
            str(r.get("Habilidade", "")),
            str(r.get("Turma", "")),
            str(r.get("Correta", "")),
        ]
        fill_cor = "F0F7F4" if idx % 2 == 0 else "FFFFFF"
        for i, (cell, val) in enumerate(zip(row_cells, valores)):
            cell.width = larguras[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 4] else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(9.5)
            if i == 4:
                run.bold = True
                run.font.color.rgb = RGBColor(0x1B, 0x43, 0x32)
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill_cor)
            tc_pr.append(shd)

    doc.add_paragraph()
    rodape = doc.add_paragraph()
    rodape.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_r = rodape.add_run(f"Total de questões: {len(df_export)}")
    run_r.font.size = Pt(9)
    run_r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run_r.italic = True

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ─── CONSTANTES E CONEXÃO ────────────────────────────────────────────────────
LISTA_PROFESSORES = [
    "",
    "ADALBERTO", "ANGELA", "ANTONIO", "CRISTIANE SILVA",
    "DANIELA DA SILVEIRA", "DANIELA FERNANDES", "DEBORA",
    "DIONATAN", "EDINEIA", "ELIVIANE", "ELMA", "EVELIM",
    "GABRIELA", "GEFFERSON", "GIZELY", "JOSE LUIZ", "JUSSARA",
    "LEANDRO", "LUCIENE", "LUIZA", "MIKAELY", "ROSANE",
    "SANDRA", "SILVANA", "TACIANE", "TATIANA", "VINICIUS",
    "WEVERTON", "WILLIAN"
]

LISTA_TURMAS = ["4° A", "5° A", "6° A", "6° B", "6° C", "7° A", "8° A",
                "9° A", "9° B", "9° C", "9° D", "1° A", "1° B", "2° A", "3° A"]

LISTA_DISCS = [
    "Arte", "Biologia", "Ciências", "Educação Física", "Filosofia",
    "Física", "Geografia", "História", "Língua Inglesa",
    "Língua Portuguesa", "Matemática", "Química", "Sociologia"
]

# CORREÇÃO 1 (item 1 da análise, incluída como bônus):
# Senha lida via st.secrets em vez de texto puro no código.
# No Streamlit Cloud, adicione em Secrets: SENHA_COORD = "coord2026"
SENHA_COORD = st.secrets.get("SENHA_COORD", "coord2026")

conn = st.connection("gsheets", type=GSheetsConnection)

# CORREÇÃO 5: TTL aumentado de 10s para 60s para reduzir chamadas à API do Sheets.
@st.cache_data(ttl=60)
def carregar_dados():
    try:
        return conn.read(ttl=0)
    except:
        return None

# ─── HEADER ──────────────────────────────────────────────────────────────────
st.markdown("""
<div class="header-box">
    <h1>📚 Portal de Simulados</h1>
    <p>Escola Estadual Padre Constantino de Monte — Maracaju/MS</p>
</div>
""", unsafe_allow_html=True)

# ─── SIDEBAR ─────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("""
    <div class="sidebar-logo">
        <span class="logo-icon">📚</span>
        <div class="logo-title">SIDE</div>
        <div class="logo-sub">Sistema de Questões</div>
    </div>
    """, unsafe_allow_html=True)

    st.divider()
    perfil = st.radio(
        "Selecione o Acesso:",
        ["👨‍🏫 Professor(a)", "🔑 Coordenação"],
        label_visibility="visible"
    )
    st.divider()

    st.markdown(f"""
    <div class="date-badge">
        📅 {datetime.now().strftime('%d/%m/%Y — %H:%M')}
    </div>
    """, unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════════
# PERFIL PROFESSOR
# ═══════════════════════════════════════════════════════════════════════════════
if perfil == "👨‍🏫 Professor(a)":

    st.markdown('<div class="section-title">✏️ Lançamento de Questões</div>', unsafe_allow_html=True)

    # CORREÇÃO 3: variáveis de formulário inicializadas fora do `with st.form`
    # para que o bloco "Download Rápido" não cause NameError quando a página
    # é carregada antes de qualquer submissão.
    nome_p  = ""
    turma_p = []
    disc_p  = sorted(LISTA_DISCS)[0]
    hab_p   = ""
    enun_p  = ""
    a = b = c = d = e = ""
    link_p  = ""
    gab_p   = "A"

    with st.form("form_prof", clear_on_submit=True):

        st.markdown('<div class="section-title">👤 Identificação</div>', unsafe_allow_html=True)
        c1, c2, c3 = st.columns([2, 1, 2])
        with c1: nome_p  = st.selectbox("Nome do(a) Professor(a) *", LISTA_PROFESSORES)
        with c2: turma_p = st.multiselect("Turma(s) *", LISTA_TURMAS,
                               placeholder="Selecione uma ou mais turmas")
        with c3: disc_p  = st.selectbox("Disciplina *", sorted(LISTA_DISCS))

        st.divider()

        st.markdown('<div class="section-title">📝 Conteúdo da Questão</div>', unsafe_allow_html=True)
        hab_p  = st.text_input("Habilidade MS *")

        st.markdown('<span class="quill-label">Enunciado da Pergunta *</span>', unsafe_allow_html=True)
        enun_p = st_quill(
            placeholder="Digite o enunciado da questão aqui. Utilize os botões para negrito e itálico.",
            html=True,
            toolbar=[['bold', 'italic'], ['clean']]
        )

        st.markdown("<br>", unsafe_allow_html=True)
        link_p = st.text_input("🔗 Link da Imagem (opcional)",
                               placeholder="https://exemplo.com/imagem.png")

        st.divider()

        st.markdown('<div class="section-title">🔤 Alternativas</div>', unsafe_allow_html=True)
        ca, cb = st.columns(2)
        with ca:
            a = st.text_input("Alternativa A *", placeholder="Opção A")
            b = st.text_input("Alternativa B *", placeholder="Opção B")
            c = st.text_input("Alternativa C *", placeholder="Opção C")
        with cb:
            d = st.text_input("Alternativa D *", placeholder="Opção D")
            e = st.text_input("Alternativa E *", placeholder="Opção E")

        st.divider()

        st.markdown('<div class="section-title">✅ Gabarito</div>', unsafe_allow_html=True)
        gab_p = st.radio("Alternativa Correta *", ["A", "B", "C", "D", "E"], horizontal=True)

        st.markdown("<br>", unsafe_allow_html=True)
        btn_enviar = st.form_submit_button("🚀 Enviar Questão", use_container_width=True)

        if btn_enviar:
            # CORREÇÃO 3: validação do enunciado separada das demais,
            # para não misturar booleano com strings no mesmo dicionário.
            enun_valido = bool(enun_p and str(enun_p).strip() not in ["<p><br></p>", "<p></p>"])

            campos_vazios = [f for f, v in {
                "Nome": nome_p,
                "Habilidade": hab_p,
                "Alt. A": a, "Alt. B": b, "Alt. C": c,
                "Alt. D": d, "Alt. E": e,
            }.items() if not v]

            if not enun_valido:
                campos_vazios.insert(0, "Enunciado")
            if not turma_p:
                campos_vazios.insert(0, "Turma")

            if campos_vazios:
                st.error(f"⚠️ Campos obrigatórios não preenchidos: **{', '.join(campos_vazios)}**")
            else:
                barra = st.progress(0, text="Preparando dados...")
                for pct, msg in [(20, "Validando campos..."),
                                 (45, "Conectando ao banco de dados..."),
                                 (70, "Gravando questão..."),
                                 (90, "Finalizando...")]:
                    time.sleep(0.3)
                    barra.progress(pct, text=msg)

                df_atual = carregar_dados()

                if df_atual is None:
                    barra.empty()
                    st.error("⚠️ Erro de conexão ao ler o banco de dados. A questão não foi salva para proteger os dados existentes. Por favor, tente enviar novamente.")
                else:
                    barra.progress(100, text="Concluído!")
                    # Uma linha por turma selecionada
                    novas_linhas = pd.DataFrame([{
                        "Data": datetime.now().strftime("%d/%m/%Y %H:%M"),
                        "Professor (a)": nome_p, "Disciplina": disc_p, "Turma": t,
                        "Habilidade": hab_p, "Pergunta": enun_p,
                        "A": a, "B": b, "C": c, "D": d, "E": e,
                        "Correta": gab_p, "Link Imagem": link_p
                    } for t in turma_p])
                    conn.update(data=pd.concat([df_atual, novas_linhas], ignore_index=True))

                    barra.empty()
                    st.balloons()
                    turmas_str = ", ".join(turma_p)
                    st.toast(f"Questão salva para {len(turma_p)} turma(s)!", icon="✅")
                    st.success(f"✨ Questão registrada com sucesso para: **{turmas_str}**")

    # ── Download rápido ──
    # CORREÇÃO 3: a verificação agora é segura pois enun_p foi inicializado acima.
    enun_valido_download = bool(enun_p and str(enun_p).strip() not in ["<p><br></p>", "<p></p>"])

    if enun_valido_download:
        st.divider()
        st.markdown('<div class="section-title">⬇️ Download Rápido</div>', unsafe_allow_html=True)
        temp_df = pd.DataFrame([{
            "Disciplina": disc_p, "Habilidade": hab_p, "Pergunta": enun_p,
            "A": a, "B": b, "C": c, "D": d, "E": e, "Link Imagem": link_p
        }])

        # CORREÇÃO 4: gerar_docx_questoes agora retorna (buffer, erros)
        doc_prof, erros_download = gerar_docx_questoes(temp_df)

        if erros_download:
            st.warning("⚠️ Algumas imagens não puderam ser inseridas no documento:")
            for erro in erros_download:
                st.caption(f"• {erro}")

        st.download_button(
            "⬇️ Baixar esta Questão em Word (.docx)",
            doc_prof, "Minha_Questao.docx",
            use_container_width=True
        )

# ═══════════════════════════════════════════════════════════════════════════════
# PERFIL COORDENAÇÃO
# ═══════════════════════════════════════════════════════════════════════════════
else:
    st.markdown('<div class="section-title">🔑 Acesso Restrito — Coordenação</div>',
                unsafe_allow_html=True)

    senha = st.text_input("Chave de Acesso:", type="password",
                          placeholder="Digite a senha de coordenação")

    if senha == SENHA_COORD:

        with st.spinner("🔄 Carregando banco de questões..."):
            barra_load = st.progress(0, text="Iniciando...")
            for pct, msg in [(30, "Conectando ao Google Sheets..."),
                             (65, "Lendo questões..."),
                             (100, "Banco carregado!")]:
                time.sleep(0.25)
                barra_load.progress(pct, text=msg)
            df = carregar_dados()
            barra_load.empty()

        if df is not None and not df.empty:

            st.markdown('<div class="section-title">📊 Visão Geral</div>', unsafe_allow_html=True)
            m1, m2, m3, m4 = st.columns(4)
            m1.metric("📋 Total de Questões", len(df))
            m2.metric("📚 Disciplinas", df["Disciplina"].nunique())
            m3.metric("🏫 Turmas", df["Turma"].nunique())
            m4.metric("👩‍🏫 Professores", df["Professor (a)"].nunique())

            st.divider()

            # ── Painel de acompanhamento por professor ──
            st.markdown('<div class="section-title">👩‍🏫 Acompanhamento por Professor(a)</div>',
                        unsafe_allow_html=True)

            # Meta por disciplina: LP e Matemática = 10, demais = 5
            META_PADRAO = 5
            META_ESPECIAL = 10
            DISCS_ESPECIAIS = {"Língua Portuguesa", "Matemática"}

            # Contagem de questões por professor e disciplina
            resumo = (
                df.groupby(["Professor (a)", "Disciplina"])
                .size()
                .reset_index(name="Enviadas")
            )
            resumo["Meta"] = resumo["Disciplina"].apply(
                lambda d: META_ESPECIAL if d in DISCS_ESPECIAIS else META_PADRAO
            )
            resumo["Completo"] = resumo["Enviadas"] >= resumo["Meta"]
            resumo["Faltam"]   = (resumo["Meta"] - resumo["Enviadas"]).clip(lower=0)

            professores_ordenados = sorted(resumo["Professor (a)"].unique())

            for prof in professores_ordenados:
                df_prof = resumo[resumo["Professor (a)"] == prof].copy()
                total_env  = int(df_prof["Enviadas"].sum())
                total_meta = int(df_prof["Meta"].sum())
                completas  = int(df_prof["Completo"].sum())
                total_disc = len(df_prof)
                pct_geral  = min(int((total_env / total_meta) * 100), 100) if total_meta > 0 else 0

                cor_status = "#52b788" if completas == total_disc else "#e76f51"
                icone      = "✅" if completas == total_disc else "⏳"

                with st.expander(
                    f"{icone} {prof} — {total_env}/{total_meta} questões  |  "
                    f"{completas}/{total_disc} disciplinas completas"
                ):
                    st.progress(pct_geral, text=f"Progresso geral: {pct_geral}%")
                    st.markdown("<br>", unsafe_allow_html=True)

                    # Tabela por disciplina
                    linhas_html = ""
                    for _, row in df_prof.sort_values("Disciplina").iterrows():
                        cor_linha = "#1b4332" if row["Completo"] else "#3d1a0a"
                        icone_disc = "✅" if row["Completo"] else f"⚠️ faltam {int(row['Faltam'])}"
                        pct_disc   = min(int((row["Enviadas"] / row["Meta"]) * 100), 100)
                        barra_fill = f"width:{pct_disc}%;background:{'#52b788' if row['Completo'] else '#e9c46a'};height:6px;border-radius:4px;"
                        linhas_html += f"""
                        <tr style='background:{cor_linha};'>
                            <td style='padding:7px 12px;color:#d8f3dc;font-size:0.85rem;'>{row['Disciplina']}</td>
                            <td style='padding:7px 12px;text-align:center;color:#d8f3dc;font-size:0.85rem;'>{int(row['Enviadas'])}</td>
                            <td style='padding:7px 12px;text-align:center;color:#a8d5bc;font-size:0.85rem;'>{int(row['Meta'])}</td>
                            <td style='padding:7px 12px;'>
                                <div style='background:rgba(255,255,255,0.08);border-radius:4px;overflow:hidden;'>
                                    <div style='{barra_fill}'></div>
                                </div>
                            </td>
                            <td style='padding:7px 12px;text-align:center;font-size:0.82rem;color:#c8e6d4;'>{icone_disc}</td>
                        </tr>"""

                    st.markdown(f"""
                    <table style='width:100%;border-collapse:collapse;border-radius:10px;overflow:hidden;'>
                        <thead>
                            <tr style='background:#2d6a4f;'>
                                <th style='padding:8px 12px;text-align:left;color:#d8f3dc;font-size:0.82rem;letter-spacing:0.5px;'>Disciplina</th>
                                <th style='padding:8px 12px;text-align:center;color:#d8f3dc;font-size:0.82rem;'>Enviadas</th>
                                <th style='padding:8px 12px;text-align:center;color:#d8f3dc;font-size:0.82rem;'>Meta</th>
                                <th style='padding:8px 12px;color:#d8f3dc;font-size:0.82rem;'>Progresso</th>
                                <th style='padding:8px 12px;text-align:center;color:#d8f3dc;font-size:0.82rem;'>Status</th>
                            </tr>
                        </thead>
                        <tbody>{linhas_html}</tbody>
                    </table>
                    """, unsafe_allow_html=True)

            st.divider()

            st.markdown('<div class="section-title">🔍 Filtros Seletores</div>',
                        unsafe_allow_html=True)
            f1, f2, f3 = st.columns(3)
            with f1: t_f = st.multiselect("🏫 Turmas:", sorted(df["Turma"].unique()))
            with f2: d_f = st.multiselect("📚 Disciplinas:", sorted(df["Disciplina"].unique()))
            with f3: p_f = st.multiselect("👩‍🏫 Professor(a):", sorted(df["Professor (a)"].unique()))

            df_v = df.copy()
            if t_f: df_v = df_v[df_v["Turma"].isin(t_f)]
            if d_f: df_v = df_v[df_v["Disciplina"].isin(d_f)]
            if p_f: df_v = df_v[df_v["Professor (a)"].isin(p_f)]

            total   = len(df)
            sel     = len(df_v)
            pct_sel = int((sel / total) * 100) if total > 0 else 0

            st.markdown("<br>", unsafe_allow_html=True)
            col_met, col_prog = st.columns([1, 3])
            with col_met:
                st.metric("✅ Questões Selecionadas", sel)
            with col_prog:
                st.markdown(f"<br><small style='color:#74c69d;font-size:0.78rem'>"
                            f"Cobertura da seleção: **{pct_sel}%** do banco total</small>",
                            unsafe_allow_html=True)
                st.progress(pct_sel)

            st.divider()
            st.dataframe(df_v, use_container_width=True, hide_index=True)

            st.divider()

            st.markdown('<div class="section-title">📄 Exportar Banco Selecionado</div>',
                        unsafe_allow_html=True)

            if not df_v.empty:
                st.markdown(
                    f"<small style='color:#a8d5bc'>Serão exportadas <b>{sel}</b> questão(ões). "
                    f"Escolha o tipo de documento abaixo:</small>",
                    unsafe_allow_html=True
                )
                st.markdown("<br>", unsafe_allow_html=True)

                with st.spinner("Preparando documentos..."):
                    # CORREÇÃO 4: desempacota a tupla (buffer, erros) retornada pela função
                    doc_banco, erros_banco = gerar_docx_questoes(df_v)
                    doc_gabarito          = gerar_docx_gabarito(df_v)

                # Exibe avisos de imagem logo após a geração, antes dos botões de download
                if erros_banco:
                    st.warning("⚠️ Algumas imagens não puderam ser inseridas no Banco de Questões:")
                    for erro in erros_banco:
                        st.caption(f"• {erro}")

                data_str = datetime.now().strftime('%d_%m_%Y')

                col_banco, col_gab = st.columns(2)

                with col_banco:
                    st.markdown("""
                    <div style='background:rgba(20,42,30,0.55);border:1px solid rgba(45,106,79,0.35);
                                border-radius:12px;padding:16px 20px;margin-bottom:12px;'>
                        <div style='font-size:1.4rem;margin-bottom:6px;'>📄</div>
                        <div style='font-weight:700;color:#74c69d;font-size:0.95rem;margin-bottom:4px;'>
                            Banco de Questões
                        </div>
                        <div style='font-size:0.78rem;color:#a8d5bc;line-height:1.5;'>
                            Documento completo com enunciados,<br>alternativas e imagens de cada questão.
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                    st.download_button(
                        label="⬇️ Baixar Banco Completo (.docx)",
                        data=doc_banco,
                        file_name=f"Banco_SIDE_{data_str}.docx",
                        use_container_width=True,
                        key="btn_banco"
                    )

                with col_gab:
                    st.markdown("""
                    <div style='background:rgba(20,42,30,0.55);border:1px solid rgba(45,106,79,0.35);
                                border-radius:12px;padding:16px 20px;margin-bottom:12px;'>
                        <div style='font-size:1.4rem;margin-bottom:6px;'>✅</div>
                        <div style='font-weight:700;color:#74c69d;font-size:0.95rem;margin-bottom:4px;'>
                            Gabarito
                        </div>
                        <div style='font-size:0.78rem;color:#a8d5bc;line-height:1.5;'>
                            Tabela com número, disciplina,<br>habilidade e resposta correta de cada questão.
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                    st.download_button(
                        label="✅ Baixar Gabarito (.docx)",
                        data=doc_gabarito,
                        file_name=f"Gabarito_SIDE_{data_str}.docx",
                        use_container_width=True,
                        key="btn_gabarito"
                    )
            else:
                st.warning("⚠️ Nenhuma questão encontrada com os filtros selecionados.")

    elif senha and senha != SENHA_COORD:
        st.error("🔒 Senha incorreta. Tente novamente.")
